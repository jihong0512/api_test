import os
import json
import yaml
import xml.etree.ElementTree as ET
import pandas as pd
from typing import List, Dict, Any, Optional
from pathlib import Path
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
import openpyxl
from pdf2image import convert_from_path
import tempfile
import shutil

from app.config import settings
from app.services.llm_service import LLMService


class DocumentParser:
    """文档解析器，支持多种格式"""
    
    def __init__(self):
        self.llm_service = LLMService()
        self.supported_formats = {
            'txt', 'json', 'yaml', 'yml', 'csv', 'xlsx', 'xls',
            'pdf', 'docx', 'md', 'jmx', 'apifox'
        }
    
    async def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """解析文档"""
        if file_type not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_type}")
        
        parser_map = {
            'txt': self._parse_txt,
            'json': self._parse_json,
            'yaml': self._parse_yaml,
            'yml': self._parse_yaml,
            'csv': self._parse_csv,
            'xlsx': self._parse_excel,
            'xls': self._parse_excel,
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'md': self._parse_markdown,
            'jmx': self._parse_jmx,
            'apifox': self._parse_apifox,
        }
        
        parser = parser_map.get(file_type)
        if not parser:
            raise ValueError(f"未实现的解析器: {file_type}")
        
        raw_data = await parser(file_path)
        
        # 使用LLM进行智能理解和标准化
        standardized_data = await self._standardize_data(raw_data, file_type)
        
        return standardized_data
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {'type': 'text', 'content': content}
    
    async def _parse_json(self, file_path: str) -> Dict[str, Any]:
        """解析JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return {'type': 'json', 'data': data}
    
    async def _parse_yaml(self, file_path: str) -> Dict[str, Any]:
        """解析YAML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        return {'type': 'yaml', 'data': data}
    
    async def _parse_csv(self, file_path: str) -> Dict[str, Any]:
        """解析CSV文件"""
        df = pd.read_csv(file_path)
        return {'type': 'csv', 'data': df.to_dict('records')}
    
    async def _parse_excel(self, file_path: str) -> Dict[str, Any]:
        """解析Excel文件"""
        workbook = openpyxl.load_workbook(file_path)
        sheets_data = {}
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            data = []
            for row in sheet.iter_rows(values_only=True):
                data.append(row)
            sheets_data[sheet_name] = data
        return {'type': 'excel', 'data': sheets_data}
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文件（使用DeepSeek-VL视觉理解）"""
        # 先尝试文本提取
        reader = PdfReader(file_path)
        text_content = ""
        for page in reader.pages:
            text_content += page.extract_text()
        
        # 对于包含图表、表格或复杂布局的PDF，使用DeepSeek-VL视觉解析
        try:
            # 将PDF转换为图片
            temp_dir = tempfile.mkdtemp()
            try:
                images = convert_from_path(file_path, dpi=200, output_folder=temp_dir)
                image_paths = []
                
                # 限制处理前5页以避免API调用过多
                for i, image in enumerate(images[:5]):
                    image_path = os.path.join(temp_dir, f"page_{i+1}.png")
                    image.save(image_path, "PNG")
                    image_paths.append(image_path)
                
                # 使用DeepSeek-VL解析
                if image_paths:
                    vision_result = await self.llm_service.parse_visual_document(
                        image_paths, 
                        document_type="pdf"
                    )
                    return {
                        'type': 'pdf',
                        'text': text_content,
                        'vision_result': vision_result
                    }
            finally:
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception as e:
            # 如果视觉解析失败，回退到文本提取
            pass
        
        return {'type': 'pdf', 'text': text_content}
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档（支持视觉理解）"""
        doc = DocxDocument(file_path)
        
        # 提取文本内容
        paragraphs = [para.text for para in doc.paragraphs]
        text_content = '\n'.join(paragraphs)
        
        # 检查是否包含表格（复杂结构，可能需要视觉理解）
        has_tables = len(doc.tables) > 0
        
        # 提取表格数据
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        result = {
            'type': 'docx',
            'content': text_content,
            'has_tables': has_tables,
            'tables': tables_data
        }
        
        # 如果包含表格，使用LLM进行结构化理解
        if has_tables:
            try:
                table_text = json.dumps(tables_data, ensure_ascii=False)
                table_prompt = f"""
请分析以下Word文档中的表格数据，提取API接口信息：

文档内容：
{text_content[:2000]}

表格数据：
{table_text[:2000]}

请识别其中的API接口信息并输出标准格式。
"""
                structured_result = await self.llm_service.chat(table_prompt)
                result['structured_data'] = structured_result
            except:
                pass
        
        return result
    
    async def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {'type': 'markdown', 'content': content}
    
    async def _parse_jmx(self, file_path: str) -> Dict[str, Any]:
        """解析JMeter JMX文件"""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        # 解析JMX结构
        test_plans = []
        for test_plan in root.findall('.//TestPlan'):
            plan_info = {
                'name': test_plan.get('testname', ''),
                'enabled': test_plan.get('enabled', 'true'),
            }
            
            # 解析HTTP采样器
            http_samples = []
            for sampler in root.findall('.//HTTPSamplerProxy'):
                sample_info = {
                    'name': sampler.get('testname', ''),
                    'domain': '',
                    'path': '',
                    'method': 'GET',
                    'port': '',
                }
                
                for element in sampler.findall('.//stringProp'):
                    prop_name = element.get('name', '')
                    prop_value = element.text or ''
                    if prop_name == 'HTTPSampler.domain':
                        sample_info['domain'] = prop_value
                    elif prop_name == 'HTTPSampler.path':
                        sample_info['path'] = prop_value
                    elif prop_name == 'HTTPSampler.method':
                        sample_info['method'] = prop_value
                    elif prop_name == 'HTTPSampler.port':
                        sample_info['port'] = prop_value
                
                http_samples.append(sample_info)
            
            plan_info['samples'] = http_samples
            test_plans.append(plan_info)
        
        return {'type': 'jmx', 'test_plans': test_plans}
    
    async def _parse_apifox(self, file_path: str) -> Dict[str, Any]:
        """解析Apifox导出的JSON格式"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Apifox格式通常包含项目、接口等信息
        return {'type': 'apifox', 'data': data}
    
    async def _standardize_data(self, raw_data: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """使用LLM标准化数据为接口信息，并初始化数据结构"""
        # 构建标准化的数据schema
        standard_schema = {
            "interfaces": [
                {
                    "name": "string",
                    "method": "GET|POST|PUT|DELETE|PATCH",
                    "url": "string",
                    "base_url": "string",
                    "headers": {},
                    "params": {},
                    "path_params": {},
                    "body": {},
                    "body_type": "json|form-data|x-www-form-urlencoded",
                    "description": "string",
                    "response_schema": {},
                    "status_codes": {}
                }
            ]
        }
        
        # 根据不同文件类型构建不同的提示
        format_instructions = self._get_format_instructions(file_type)
        
        prompt = f"""
请将以下{file_type}格式的文档解析为标准的API接口信息。

{format_instructions}

请严格按照以下JSON Schema输出标准化的接口信息列表：
{json.dumps(standard_schema, ensure_ascii=False, indent=2)}

原始数据：
{json.dumps(raw_data, ensure_ascii=False, indent=2)[:8000]}

要求：
1. 提取所有API接口信息
2. 确保URL、参数、请求体等信息完整准确
3. 对于复杂嵌套结构，请展开为扁平化格式
4. 如果原始数据中包含视觉解析结果（vision_result），请优先使用视觉解析的结果
5. 输出必须是有效的JSON格式，只输出JSON，不要包含其他文字

请输出标准化的接口信息列表（JSON格式）：
"""
        
        try:
            result = await self.llm_service.chat(prompt, temperature=0.3, max_tokens=4000)
            
            # 清理结果，提取JSON部分
            import re
            json_match = re.search(r'\{[\s\S]*\}', result)
            if json_match:
                json_str = json_match.group()
                standardized = json.loads(json_str)
            else:
                standardized = json.loads(result)
            
            # 验证和补充数据结构
            standardized = self._validate_and_complete_structure(standardized)
            
            return standardized
        except json.JSONDecodeError as e:
            # 如果JSON解析失败，尝试从文本中提取
            return self._fallback_extract(raw_data, result if 'result' in locals() else str(raw_data))
        except Exception as e:
            # 如果LLM调用失败，返回原始数据
            return {"interfaces": [], "error": str(e), "raw_data": raw_data}
    
    def _get_format_instructions(self, file_type: str) -> str:
        """获取不同格式的解析指令"""
        instructions_map = {
            "apifox": """
Apifox格式说明：
- 通常包含projects、apis等字段
- api.method 表示HTTP方法
- api.path 表示接口路径
- api.request 包含请求参数、请求头、请求体
- api.response 包含响应信息
""",
            "jmx": """
JMeter JMX格式说明：
- HTTPSamplerProxy 表示HTTP请求采样器
- HTTPSampler.method 表示HTTP方法
- HTTPSampler.domain + HTTPSampler.path 组成完整URL
- HTTPSampler.port 表示端口号
- 可能包含参数、请求头等信息
""",
            "json": """
JSON格式说明：
- 可能直接是OpenAPI/Swagger格式
- 也可能包含apis、endpoints等字段
- 识别path、method、parameters、requestBody等字段
""",
            "yaml": """
YAML格式说明：
- 通常是OpenAPI/Swagger YAML格式
- paths下包含各个接口路径
- 每个路径下有method（get/post等）
- 包含parameters、requestBody、responses等
""",
            "pdf": """
PDF格式说明：
- 可能包含接口文档、API说明等
- 注意提取表格中的接口信息
- 识别URL、方法、参数等关键信息
- 如果包含视觉解析结果（vision_result），请使用视觉解析的结果
""",
            "docx": """
Word格式说明：
- 可能包含接口文档、表格等
- 表格中可能包含接口列表
- 识别method、url、参数等信息
""",
        }
        return instructions_map.get(file_type, "请提取其中的API接口信息。")
    
    def _validate_and_complete_structure(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """验证和补充数据结构，确保完整性"""
        if not isinstance(data, dict):
            return {"interfaces": []}
        
        interfaces = data.get("interfaces", [])
        if not isinstance(interfaces, list):
            interfaces = []
        
        # 标准化每个接口的数据结构
        standardized_interfaces = []
        for iface in interfaces:
            if not isinstance(iface, dict):
                continue
            
            standardized = {
                "name": iface.get("name", ""),
                "method": iface.get("method", "GET").upper(),
                "url": iface.get("url", ""),
                "base_url": iface.get("base_url", ""),
                "headers": iface.get("headers", {}),
                "params": iface.get("params", {}),
                "path_params": iface.get("path_params", {}),
                "body": iface.get("body", {}),
                "body_type": iface.get("body_type", "json"),
                "description": iface.get("description", ""),
                "response_schema": iface.get("response_schema", {}),
                "status_codes": iface.get("status_codes", {})
            }
            
            # 确保URL完整
            if standardized["url"] and standardized["base_url"]:
                if not standardized["url"].startswith("http"):
                    standardized["url"] = standardized["base_url"].rstrip("/") + "/" + standardized["url"].lstrip("/")
            
            standardized_interfaces.append(standardized)
        
        return {"interfaces": standardized_interfaces}
    
    def _fallback_extract(self, raw_data: Dict[str, Any], text: str) -> Dict[str, Any]:
        """后备提取方法，从原始数据或文本中提取接口信息"""
        interfaces = []
        
        # 尝试从不同类型的原始数据中提取
        if raw_data.get("type") == "apifox" and "data" in raw_data:
            interfaces = self._extract_from_apifox(raw_data["data"])
        elif raw_data.get("type") == "jmx" and "test_plans" in raw_data:
            interfaces = self._extract_from_jmx(raw_data["test_plans"])
        elif raw_data.get("type") in ["json", "yaml"] and "data" in raw_data:
            interfaces = self._extract_from_openapi(raw_data["data"])
        
        return {"interfaces": interfaces}
    
    def _extract_from_apifox(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从Apifox数据中提取接口"""
        interfaces = []
        # Apifox格式解析逻辑
        if isinstance(data, dict):
            apis = data.get("apis", [])
            for api in apis:
                interfaces.append({
                    "name": api.get("name", ""),
                    "method": api.get("method", "GET"),
                    "url": api.get("path", ""),
                    "description": api.get("description", "")
                })
        return interfaces
    
    def _extract_from_jmx(self, test_plans: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """从JMX数据中提取接口"""
        interfaces = []
        for plan in test_plans:
            for sample in plan.get("samples", []):
                url = f"{sample.get('domain', '')}{sample.get('path', '')}"
                interfaces.append({
                    "name": sample.get("name", ""),
                    "method": sample.get("method", "GET"),
                    "url": url,
                    "description": f"JMeter测试采样器"
                })
        return interfaces
    
    def _extract_from_openapi(self, data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从OpenAPI/Swagger数据中提取接口"""
        interfaces = []
        if isinstance(data, dict) and "paths" in data:
            base_url = data.get("servers", [{}])[0].get("url", "") if data.get("servers") else ""
            for path, methods in data["paths"].items():
                for method, details in methods.items():
                    if method.lower() in ["get", "post", "put", "delete", "patch"]:
                        interfaces.append({
                            "name": details.get("summary", details.get("operationId", path)),
                            "method": method.upper(),
                            "url": path,
                            "base_url": base_url,
                            "description": details.get("description", ""),
                            "params": {param["name"]: param for param in details.get("parameters", [])},
                            "body": details.get("requestBody", {}).get("content", {})
                        })
        return interfaces
    
    def extract_api_interfaces(self, standardized_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从标准化数据中提取接口列表，确保数据结构完整"""
        interfaces = []
        
        if isinstance(standardized_data, list):
            # 如果是列表，直接使用
            interfaces = standardized_data
        elif isinstance(standardized_data, dict):
            # 根据不同的数据结构提取接口
            if 'interfaces' in standardized_data:
                interfaces = standardized_data['interfaces']
            elif 'apis' in standardized_data:
                interfaces = standardized_data['apis']
            else:
                # 尝试直接解析为单个接口
                if 'name' in standardized_data and 'url' in standardized_data:
                    interfaces = [standardized_data]
        
        # 验证和清理接口数据
        validated_interfaces = []
        for iface in interfaces:
            if isinstance(iface, dict) and iface.get("url"):
                # 确保必要字段存在
                validated = {
                    "name": iface.get("name", iface.get("url", "")),
                    "method": iface.get("method", "GET").upper(),
                    "url": iface.get("url", ""),
                    "base_url": iface.get("base_url", ""),
                    "headers": iface.get("headers", {}),
                    "params": iface.get("params", {}),
                    "path_params": iface.get("path_params", {}),
                    "body": iface.get("body", {}),
                    "body_type": iface.get("body_type", "json"),
                    "description": iface.get("description", ""),
                    "response_schema": iface.get("response_schema", {}),
                    "status_codes": iface.get("status_codes", {})
                }
                validated_interfaces.append(validated)
        
        return validated_interfaces


