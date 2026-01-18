import os
import json
import base64
import yaml
import xml.etree.ElementTree as ET
import pandas as pd
import re
from typing import List, Dict, Any, Optional
from pathlib import Path
# 延迟导入docx，避免在不需要时导入失败
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
from pdf2image import convert_from_path
import tempfile
import shutil
import requests
from PIL import Image

from app.config import settings


class EnhancedDocumentParser:
    """增强的文档解析器：智能选择解析方式"""
    
    def __init__(self):
        self.qwen_api_key = settings.QWEN_API_KEY
        self.qwen_base_url = settings.QWEN_BASE_URL
        self.qwen_model = settings.QWEN_MODEL
        self.supported_formats = {
            'pdf', 'docx', 'doc', 'xlsx', 'xls', 'xmind', 'csv', 
            'json', 'yaml', 'yml', 'txt', 'md', 'xml',
            'postman', 'apifox', 'jmx', 'swagger'
        }
    
    async def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """解析文档，返回分类结果 - 统一使用库解析，不使用qwen-vl-plus"""
        if file_type not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_type}")
        
        # 所有文件类型统一使用库解析
        return await self._parse_with_library(file_path, file_type)
    
    # 以下方法已废弃，不再使用qwen-vl-plus解析
    # async def _detect_complex_content(self, file_path: str, file_type: str) -> bool:
        """检测文件是否包含图片、表格、公式等复杂内容"""
        try:
            if file_type == 'pdf':
                # 使用pdfplumber检测PDF中的表格和图片
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages[:5]:  # 检查前5页
                        # 检测表格
                        tables = page.extract_tables()
                        if tables:
                            return True
                        
                        # 检测图片（通过图像对象）
                        if hasattr(page, 'images') and page.images:
                            return True
                        
                        # 检测公式（通过文本中的特殊字符）
                        text = page.extract_text() or ""
                        if any(char in text for char in ['∑', '∫', '√', '∆', 'α', 'β', 'γ', 'π', '=', '≤', '≥']):
                            return True
                
                # 如果没有检测到，返回False
                return False
            
            elif file_type in {'docx', 'doc'}:
                # 检测Word中的表格、图片
                try:
                    if DocxDocument is None:
                        return False  # 如果docx未安装，返回False，不使用qwen-vl
                    doc = DocxDocument(file_path)
                    
                    # 检测表格
                    if doc.tables:
                        return True
                    
                    # 检测图片（通过关系）
                    # python-docx可以通过part.rels检测图片
                    if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
                        for rel in doc.part.related_parts.values():
                            if hasattr(rel, 'target_ref') and any(img_ext in str(rel.target_ref) for img_ext in ['.png', '.jpg', '.jpeg', '.gif']):
                                return True
                    
                    # 检测公式（通过文本）
                    for para in doc.paragraphs:
                        text = para.text
                        if any(char in text for char in ['∑', '∫', '√', '∆', 'α', 'β', 'γ', 'π', '=', '≤', '≥']):
                            return True
                    
                    return False
                except Exception:
                    # 如果检测失败，默认使用qwen-vl（安全起见）
                    return True
            
            return False
        except Exception as e:
            # 检测失败时，默认使用qwen-vl
            print(f"检测复杂内容失败: {e}，将使用qwen-vl解析")
            return True
    
    async def _parse_with_qwen(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """使用qwen-vl-plus解析文件（适用于包含图片、表格、公式的文件）"""
        try:
            # 准备图片
            if file_type in {'jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp'}:
                images = [file_path] if os.path.exists(file_path) else []
            elif file_type == 'pdf':
                images = await self._pdf_to_images(file_path)
            elif file_type in {'docx', 'doc'}:
                # Word需要转换为图片
                images = await self._docx_to_images(file_path)
            else:
                images = []
            
            if not images:
                raise ValueError(f"无法将{file_type}文件转换为图片")
            
            # 调用qwen-vl-plus API
            result = await self._call_qwen_vl_api(images, file_type)
            
            # 分类解析结果
            classified_result = self._classify_result(result, file_type)
            
            return classified_result
            
        except Exception as e:
            raise Exception(f"qwen-vl-plus解析失败: {str(e)}")
    
    async def _parse_with_library(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """使用库解析文件（所有文件类型统一使用库解析）"""
        parser_map = {
            'pdf': self._parse_pdf_library,
            'docx': self._parse_docx_library,
            'doc': self._parse_docx_library,
            'xlsx': self._parse_excel_library,
            'xls': self._parse_excel_library,
            'csv': self._parse_csv_library,
            'xmind': self._parse_xmind_library,
            'json': self._parse_json,
            'yaml': self._parse_yaml,
            'yml': self._parse_yaml,
            'txt': self._parse_txt,
            'md': self._parse_markdown,
            'xml': self._parse_xml,
            'postman': self._parse_postman,
            'apifox': self._parse_apifox,
            'jmx': self._parse_jmx,
            'swagger': self._parse_swagger,
        }
        
        parser = parser_map.get(file_type)
        if not parser:
            raise ValueError(f"未实现的解析器: {file_type}")
        
        raw_data = await parser(file_path)
        
        # 转换为统一格式
        return self._convert_to_classified_format(raw_data, file_type)
    
    async def _parse_pdf_library(self, file_path: str) -> Dict[str, Any]:
        """使用pdfplumber解析PDF"""
        if pdfplumber is None:
            raise ImportError("pdfplumber模块未安装，无法解析PDF文档。请运行: pip install pdfplumber")
        text_content = []
        tables = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(text)
                    
                    # 提取表格
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables, 1):
                        if table:
                            headers = table[0] if len(table) > 0 else []
                            rows = table[1:] if len(table) > 1 else []
                            tables.append({
                                'table_index': table_idx,
                                'page': page_num,
                                'headers': headers,
                                'rows': rows
                            })
        except Exception as e:
            print(f"pdfplumber解析失败: {e}")
            # 降级使用PyPDF2
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_content = [page.extract_text() for page in reader.pages if page.extract_text()]
        
        return {
            'type': 'pdf',
            'text_content': text_content,
            'tables': tables,
            'metadata': {'parser': 'pdfplumber'}
        }
    
    async def _parse_docx_library(self, file_path: str) -> Dict[str, Any]:
        """使用python-docx解析Word文档"""
        if DocxDocument is None:
            raise ImportError("python-docx模块未安装，无法解析Word文档。请运行: pip install python-docx")
        doc = DocxDocument(file_path)
        
        # 提取文本
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # 提取表格
        tables = []
        for table_idx, table in enumerate(doc.tables, 1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                headers = table_data[0] if len(table_data) > 0 else []
                rows = table_data[1:] if len(table_data) > 1 else []
                tables.append({
                    'table_index': table_idx,
                    'headers': headers,
                    'rows': rows
                })
        
        return {
            'type': 'docx',
            'text_content': text_content,
            'tables': tables,
            'metadata': {'parser': 'python-docx'}
        }
    
    async def _parse_excel_library(self, file_path: str) -> Dict[str, Any]:
        """使用openpyxl解析Excel"""
        if openpyxl is None:
            raise ImportError("openpyxl模块未安装，无法解析Excel文档。请运行: pip install openpyxl")
        workbook = openpyxl.load_workbook(file_path)
        tables = []
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                sheet_data.append(row_data)
            
            if sheet_data:
                # 将每行作为文本内容
                for row in sheet_data:
                    text_content.append(' | '.join(row))
                
                # 第一个非空行作为表头
                headers = sheet_data[0] if sheet_data else []
                rows = sheet_data[1:] if len(sheet_data) > 1 else []
                
                tables.append({
                    'table_index': len(tables) + 1,
                    'sheet_name': sheet_name,
                    'headers': headers,
                    'rows': rows
                })
        
        return {
            'type': 'excel',
            'text_content': text_content,
            'tables': tables,
            'metadata': {'parser': 'openpyxl', 'sheets': workbook.sheetnames}
        }
    
    async def _parse_csv_library(self, file_path: str) -> Dict[str, Any]:
        """使用csv库解析CSV文件"""
        try:
            import csv
            text_content = []
            rows = []
            headers = []
            
            with open(file_path, 'r', encoding='utf-8') as f:
                # 尝试自动检测分隔符
                sample = f.read(1024)
                f.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(f, delimiter=delimiter)
                for row_idx, row in enumerate(reader):
                    if row_idx == 0:
                        headers = row
                    else:
                        rows.append(row)
                        text_content.append(' | '.join([str(val) for val in row]))
            
            tables = [{
                'table_index': 1,
                'headers': headers,
                'rows': rows
            }]
            
            return {
                'type': 'csv',
                'text_content': text_content,
                'tables': tables,
                'metadata': {'parser': 'csv', 'row_count': len(rows), 'delimiter': delimiter}
            }
        except Exception as e:
            raise Exception(f"CSV解析失败: {str(e)}")
    
    async def _parse_xmind_library(self, file_path: str) -> Dict[str, Any]:
        """解析XMind文件"""
        try:
            import xmindparser
            workbook = xmindparser.xmind_to_dict(file_path)
            
            # 转换为文本
            text_content = [json.dumps(workbook, ensure_ascii=False)]
            
            return {
                'type': 'xmind',
                'text_content': text_content,
                'tables': [],
                'metadata': {'parser': 'xmindparser', 'raw_data': workbook}
            }
        except Exception as e:
            raise Exception(f"XMind解析失败: {str(e)}")
    
    def _convert_to_classified_format(self, raw_data: Dict[str, Any], file_type: str) -> Dict[str, Any]:
        """将库解析结果转换为统一格式"""
        # 处理text_content字段，可能是列表或字符串
        text_content = raw_data.get("text_content", [])
        if isinstance(text_content, str):
            text_content = [text_content] if text_content.strip() else []
        elif not isinstance(text_content, list):
            text_content = []
        
        # 处理content字段（txt, json, yaml等可能使用content字段）
        if not text_content and "content" in raw_data:
            content = raw_data["content"]
            if isinstance(content, str):
                text_content = [content] if content.strip() else []
            elif isinstance(content, list):
                text_content = content
        
        # JSON文件解析可能直接返回text_content
        if not text_content and "text_content" in raw_data:
            text_content = raw_data["text_content"]
            if isinstance(text_content, str):
                text_content = [text_content] if text_content.strip() else []
        
        result = {
            "text": text_content,
            "images": [],
            "tables": raw_data.get("tables", []),
            "formulas": [],
            "metadata": raw_data.get("metadata", {"file_type": file_type}),
            "raw_result": raw_data
        }
        
        # Swagger/Postman/Apifox等格式会直接返回interfaces，保留原始数据和interfaces
        if file_type in ["postman", "apifox", "jmx", "swagger"]:
            if "interfaces" in raw_data:
                result["interfaces"] = raw_data["interfaces"]
        
        # JSON文件如果包含interfaces，也要保留（JSON请求日志格式）
        if file_type == "json" and "interfaces" in raw_data:
            result["interfaces"] = raw_data["interfaces"]
        
        # YAML文件如果被检测为OpenAPI/Swagger格式，也会包含interfaces（_parse_yaml会自动检测）
        # 检查raw_data的type字段，如果是swagger，说明是OpenAPI格式
        if file_type in ["yaml", "yml"]:
            # 如果raw_data的type是swagger，说明_parse_yaml检测到了OpenAPI格式
            if raw_data.get("type") == "swagger" and "interfaces" in raw_data:
                result["interfaces"] = raw_data["interfaces"]
            # 或者直接检查是否有interfaces字段（更通用）
            elif "interfaces" in raw_data:
                result["interfaces"] = raw_data["interfaces"]
        
        return result
    
    async def _call_qwen_vl_api(self, image_paths: List[str], file_type: str) -> str:
        """调用qwen-vl-plus API"""
        prompt = f"""
请解析这个{file_type}文件的内容，并按照以下要求分类提取：

1. **文本内容**：提取所有文本信息，包括段落、标题、描述等
2. **图片内容**：识别和描述图片中的内容
3. **表格内容**：提取所有表格数据，包括表头、行、列数据
4. **公式内容**：识别并提取数学公式、化学公式等公式内容

请以JSON格式输出，格式如下：
{{
    "text_content": ["文本段落1", "文本段落2", ...],
    "image_descriptions": ["图片1描述", "图片2描述", ...],
    "tables": [
        {{
            "table_index": 1,
            "headers": ["列1", "列2", ...],
            "rows": [
                ["数据1", "数据2", ...],
                ...
            ]
        }}
    ],
    "formulas": ["公式1", "公式2", ...],
    "metadata": {{
        "file_type": "{file_type}",
        "total_pages": 数量,
        "has_tables": true/false,
        "has_formulas": true/false
    }}
}}

请仔细分析文件内容，确保提取的信息准确完整。
"""
        
        try:
            content = [{"type": "text", "text": prompt}]
            
            # 添加图片
            for image_path in image_paths[:10]:  # 限制最多10张图片
                base64_image = self._encode_image(image_path)
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{base64_image}"
                    }
                })
            
            headers = {
                "Authorization": f"Bearer {self.qwen_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": self.qwen_model,
                "messages": [
                    {
                        "role": "user",
                        "content": content
                    }
                ],
                "max_tokens": 8000
            }
            
            response = requests.post(
                f"{self.qwen_base_url}/chat/completions",
                headers=headers,
                json=payload,
                timeout=120
            )
            
            if response.status_code != 200:
                raise Exception(f"API调用失败: {response.status_code}, {response.text}")
            
            result = response.json()
            return result['choices'][0]['message']['content']
            
        except Exception as e:
            raise Exception(f"调用qwen-vl-plus API失败: {str(e)}")
    
    def _classify_result(self, qwen_result: str, file_type: str) -> Dict[str, Any]:
        """分类解析结果"""
        try:
            import re
            json_match = re.search(r'\{[\s\S]*\}', qwen_result)
            if json_match:
                parsed = json.loads(json_match.group())
            else:
                parsed = json.loads(qwen_result)
            
            result = {
                "text": parsed.get("text_content", []),
                "images": parsed.get("image_descriptions", []),
                "tables": parsed.get("tables", []),
                "formulas": parsed.get("formulas", []),
                "metadata": parsed.get("metadata", {}),
                "raw_result": qwen_result
            }
            
            return result
            
        except json.JSONDecodeError:
            return {
                "text": [qwen_result],
                "images": [],
                "tables": [],
                "formulas": [],
                "metadata": {"file_type": file_type},
                "raw_result": qwen_result
            }
    
    async def _pdf_to_images(self, file_path: str) -> List[str]:
        """PDF转换为图片"""
        temp_dir = tempfile.mkdtemp()
        try:
            images = convert_from_path(file_path, dpi=200, output_folder=temp_dir)
            image_paths = []
            for i, image in enumerate(images[:20]):
                image_path = os.path.join(temp_dir, f"page_{i+1}.png")
                image.save(image_path, "PNG")
                image_paths.append(image_path)
            return image_paths
        except Exception as e:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)
            raise Exception(f"PDF转图片失败: {str(e)}")
    
    async def _docx_to_images(self, file_path: str) -> List[str]:
        """Word文档转换为图片"""
        # 简化处理：对于需要视觉解析的Word，可以通过pdf中转
        # 这里暂时返回空，如果需要完整实现可以使用unoconv或LibreOffice
        return []
    
    def _encode_image(self, image_path: str) -> str:
        """将图片编码为base64"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    async def _parse_json(self, file_path: str) -> Dict[str, Any]:
        """解析JSON文件，支持ks_all_interface.json格式"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # 检查是否是API请求日志格式（包含url、request_body、response_body等字段）
        interfaces = []
        
        # 处理ks_all_interface.json格式（可能是对象，包含interfaces数组）
        if isinstance(data, dict):
            # 如果包含interfaces字段，直接使用
            if 'interfaces' in data and isinstance(data['interfaces'], list):
                data = data['interfaces']
            # 如果包含data字段且是数组
            elif 'data' in data and isinstance(data['data'], list):
                data = data['data']
            # 如果包含items字段且是数组
            elif 'items' in data and isinstance(data['items'], list):
                data = data['items']
        
        if isinstance(data, list):
            # 处理请求日志数组格式
            for item in data:
                if isinstance(item, dict):
                    # 提取接口标题（title字段）
                    title = item.get('title', '')
                    
                    # 提取URL（优先使用url字段，然后是path字段）
                    url_from_item = item.get('url', '') or item.get('path', '') or item.get('endpoint', '')
                    
                    # 提取base_url和path
                    base_url = item.get('base_url', '') or item.get('baseURL', '')
                    
                    # 如果url_from_item是完整URL，解析它
                    if url_from_item and (url_from_item.startswith("http://") or url_from_item.startswith("https://")):
                        # 已经是完整URL
                        url = url_from_item
                        from urllib.parse import urlparse
                        parsed = urlparse(url)
                        base_url = f"{parsed.scheme}://{parsed.netloc}"
                        path = parsed.path + ('?' + parsed.query if parsed.query else '')
                    elif base_url and url_from_item:
                        # 有base_url和path，组合成完整URL
                        path = url_from_item
                        url = base_url.rstrip('/') + '/' + path.lstrip('/')
                    elif url_from_item:
                        # 只有path或url，使用它
                        url = url_from_item
                        path = url_from_item
                        # 尝试从URL中提取base_url
                        if url.startswith("http://") or url.startswith("https://"):
                            from urllib.parse import urlparse
                            parsed = urlparse(url)
                            base_url = f"{parsed.scheme}://{parsed.netloc}"
                            path = parsed.path + ('?' + parsed.query if parsed.query else '')
                    else:
                        # 没有URL信息，跳过
                        continue
                    
                    # 提取HTTP方法（如果没有method字段，根据是否有request_body判断，默认POST）
                    method = item.get('method', 'POST').upper()  # 默认POST，支持GET/POST/PUT/DELETE等
                    
                    # 提取请求体（优先使用request_body，然后是body字段）
                    request_body = item.get('request_body', item.get('body', item.get('request', {})))
                    if isinstance(request_body, str):
                        try:
                            request_body = json.loads(request_body)
                        except:
                            request_body = {}
                    
                    # 从顶层或request_body中提取service字段
                    service = item.get('service', '')
                    if not service and isinstance(request_body, dict):
                        service = request_body.get('service', '')
                    # service字段可能大小写不同，统一处理
                    if service:
                        service = service.strip()
                    
                    # 接口名称：优先使用title，其次是name，再是service，最后是path的最后一部分
                    name = title or item.get('name') or service or item.get('interface_name') or path.split('/')[-1].split('?')[0] or 'API接口'
                    
                    # 提取请求头
                    headers = item.get('headers', {})
                    if isinstance(headers, str):
                        try:
                            headers = json.loads(headers)
                        except:
                            headers = {}
                    
                    # 提取URL参数
                    params = item.get('params', item.get('query', {}))
                    if isinstance(params, str):
                        try:
                            params = json.loads(params)
                        except:
                            params = {}
                    
                    # 提取响应信息（预留字段，即使没有数据也保留）
                    response_body = item.get('response_body', item.get('response', item.get('responseBody', {})))
                    if isinstance(response_body, str):
                        try:
                            response_body = json.loads(response_body)
                        except:
                            response_body = {}
                    # 如果没有响应体，设为空对象
                    if not response_body:
                        response_body = {}
                    
                    # 提取响应头（预留字段）
                    response_headers = item.get('response_headers', item.get('responseHeaders', item.get('response_header', {})))
                    if isinstance(response_headers, str):
                        try:
                            response_headers = json.loads(response_headers)
                        except:
                            response_headers = {}
                    # 如果没有响应头，设为空对象
                    if not response_headers:
                        response_headers = {}
                    
                    response_schema = item.get('response_schema', item.get('responseSchema', {}))
                    if isinstance(response_schema, str):
                        try:
                            response_schema = json.loads(response_schema)
                        except:
                            response_schema = {}
                    # 如果没有响应schema，设为空对象
                    if not response_schema:
                        response_schema = {}
                    
                    status_code = item.get('status_code', item.get('statusCode', 200))
                    description = item.get('description', item.get('desc', ''))
                    
                    # 提取版本信息
                    version = item.get('version', '')
                    
                    interfaces.append({
                        'name': name,
                        'method': method,
                        'url': url,
                        'base_url': base_url,
                        'path': path,
                        'request_body': request_body,
                        'response_body': response_body,
                        'response_headers': response_headers,  # 添加响应头字段
                        'response_schema': response_schema,
                        'headers': headers,
                        'params': params,
                        'status_code': status_code,
                        'service': service,
                        'description': description or (f'{name} - {service}' if service else name),
                        'version': version,  # 添加版本字段
                        'tags': item.get('tags', []) if isinstance(item.get('tags'), list) else []
                    })
        
        result = {
            'type': 'json',
            'data': data,
            'metadata': {'parser': 'json', 'is_request_log': isinstance(data, list) and len(data) > 0 and isinstance(data[0], dict) and 'url' in data[0]}
        }
        
        # 如果有接口，添加到结果中
        if interfaces:
            result['interfaces'] = interfaces
        
        # 转换为文本格式（用于全文搜索）
        if isinstance(data, (dict, list)):
            text_content = [json.dumps(data, ensure_ascii=False, indent=2)]
        else:
            text_content = [str(data)]
        result['text_content'] = text_content
        
        return result
    
    async def _parse_yaml(self, file_path: str) -> Dict[str, Any]:
        """解析YAML文件 - 使用PyYAML库，如果是OpenAPI/Swagger格式则自动识别"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        # 检测是否是OpenAPI/Swagger格式
        if isinstance(data, dict):
            # 检查是否包含OpenAPI/Swagger的特征字段
            if 'openapi' in data or 'swagger' in data:
                # 如果是OpenAPI/Swagger格式，使用swagger解析器处理
                interfaces = self._extract_swagger_interfaces(data)
                return {
                    'type': 'swagger',
                    'interfaces': interfaces,
                    'swagger_version': data.get('swagger') or data.get('openapi'),
                    'info': data.get('info', {}),
                    'servers': data.get('servers', []),
                    'paths': data.get('paths', {}),
                    'raw_data': data
                }
        
        # 普通YAML文件
        return {'type': 'yaml', 'data': data, 'metadata': {'parser': 'PyYAML'}}
    
    async def _parse_xml(self, file_path: str) -> Dict[str, Any]:
        """解析XML文件"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # 转换为字典格式（简单处理）
            def element_to_dict(elem):
                result = {
                    'tag': elem.tag,
                    'text': elem.text.strip() if elem.text and elem.text.strip() else None,
                    'attrib': elem.attrib if elem.attrib else None,
                    'children': [element_to_dict(child) for child in elem] if len(elem) > 0 else None
                }
                return result
            
            data = element_to_dict(root)
            
            # 提取文本内容
            text_content = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_content.append(elem.text.strip())
            
            return {
                'type': 'xml',
                'data': data,
                'text_content': text_content,
                'metadata': {'parser': 'xml.etree.ElementTree'}
            }
        except Exception as e:
            raise Exception(f"XML解析失败: {str(e)}")
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            # 如果UTF-8失败，尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    content = f.read()
            except:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
        
        return {'type': 'text', 'content': content, 'metadata': {'parser': 'plain'}}
    
    async def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """解析Markdown文件 - 使用python-markdown库"""
        try:
            import markdown
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            # 可以将markdown转换为HTML，也可以直接保留原始内容
            html_content = markdown.markdown(content, extensions=['extra', 'codehilite'])
            return {
                'type': 'markdown',
                'content': content,
                'html_content': html_content,
                'metadata': {'parser': 'python-markdown'}
            }
        except ImportError:
            # 如果markdown库不存在，降级为纯文本读取
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {'type': 'markdown', 'content': content, 'metadata': {'parser': 'plain'}}
    
    async def _parse_postman(self, file_path: str) -> Dict[str, Any]:
        """解析Postman导出的JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        interfaces = []
        if 'item' in data:
            def extract_items(items):
                for item in items:
                    if 'request' in item:
                        request = item.get('request', {})
                        url_obj = request.get('url', {})
                        url = url_obj.get('raw', url_obj.get('host', [''])[0] if isinstance(url_obj.get('host'), list) else '')
                        
                        interfaces.append({
                            'name': item.get('name', ''),
                            'method': request.get('method', 'GET'),
                            'url': url,
                            'headers': {h.get('key'): h.get('value') for h in request.get('header', [])},
                            'params': {p.get('key'): p.get('value') for p in url_obj.get('query', [])},
                            'body': request.get('body', {}),
                            'description': request.get('description', '')
                        })
                    if 'item' in item:
                        extract_items(item.get('item', []))
            
            extract_items(data.get('item', []))
        
        return {'type': 'postman', 'interfaces': interfaces, 'raw_data': data}
    
    async def _parse_apifox(self, file_path: str) -> Dict[str, Any]:
        """解析Apifox导出的JSON文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        interfaces = []
        if isinstance(data, dict):
            apis = data.get('apis', data.get('data', {}).get('apis', []))
            for api in apis:
                interfaces.append({
                    'name': api.get('name', ''),
                    'method': api.get('method', 'GET'),
                    'url': api.get('path', api.get('url', '')),
                    'headers': api.get('headers', {}),
                    'params': api.get('params', {}),
                    'body': api.get('requestBody', api.get('body', {})),
                    'description': api.get('description', ''),
                    'response_schema': api.get('response', {})
                })
        
        return {'type': 'apifox', 'interfaces': interfaces, 'raw_data': data}
    
    async def _parse_jmx(self, file_path: str) -> Dict[str, Any]:
        """解析JMeter JMX文件"""
        tree = ET.parse(file_path)
        root = tree.getroot()
        
        test_plans = []
        for test_plan in root.findall('.//TestPlan'):
            plan_info = {
                'name': test_plan.get('testname', ''),
                'enabled': test_plan.get('enabled', 'true'),
            }
            
            http_samples = []
            for sampler in root.findall('.//HTTPSamplerProxy'):
                sample_info = {
                    'name': sampler.get('testname', ''),
                    'domain': '',
                    'path': '',
                    'method': 'GET',
                    'port': '',
                }
                
                for element in sampler.findall('.//stringProp'):
                    prop_name = element.get('name', '')
                    prop_value = element.text or ''
                    if prop_name == 'HTTPSampler.domain':
                        sample_info['domain'] = prop_value
                    elif prop_name == 'HTTPSampler.path':
                        sample_info['path'] = prop_value
                    elif prop_name == 'HTTPSampler.method':
                        sample_info['method'] = prop_value
                    elif prop_name == 'HTTPSampler.port':
                        sample_info['port'] = prop_value
                
                http_samples.append(sample_info)
            
            plan_info['samples'] = http_samples
            test_plans.append(plan_info)
        
        interfaces = []
        for plan in test_plans:
            for sample in plan.get('samples', []):
                url = f"{sample.get('domain', '')}{sample.get('path', '')}"
                interfaces.append({
                    'name': sample.get('name', ''),
                    'method': sample.get('method', 'GET'),
                    'url': url,
                    'description': 'JMeter测试采样器'
                })
        
        return {'type': 'jmx', 'test_plans': test_plans, 'interfaces': interfaces}
    
    async def _parse_swagger(self, file_path: str) -> Dict[str, Any]:
        """解析Swagger/OpenAPI文件（支持本地文件和在线URL）"""
        try:
            # 判断是URL还是本地文件路径
            if file_path.startswith('http://') or file_path.startswith('https://'):
                # 在线Swagger URL
                response = requests.get(file_path, timeout=30)
                response.raise_for_status()
                
                # 根据Content-Type判断格式
                content_type = response.headers.get('Content-Type', '').lower()
                if 'yaml' in content_type or 'yml' in content_type:
                    swagger_data = yaml.safe_load(response.text)
                else:
                    swagger_data = response.json()
            else:
                # 本地文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    # 尝试YAML格式
                    try:
                        swagger_data = yaml.safe_load(f)
                    except yaml.YAMLError:
                        # 如果不是YAML，尝试JSON
                        f.seek(0)
                        swagger_data = json.load(f)
            
            # 提取API接口
            interfaces = self._extract_swagger_interfaces(swagger_data)
            
            return {
                'type': 'swagger',
                'interfaces': interfaces,
                'swagger_version': swagger_data.get('swagger') or swagger_data.get('openapi'),
                'info': swagger_data.get('info', {}),
                'servers': swagger_data.get('servers', []),
                'paths': swagger_data.get('paths', {}),
                'raw_data': swagger_data
            }
        except Exception as e:
            raise Exception(f"Swagger解析失败: {str(e)}")
    
    def _extract_swagger_interfaces(self, swagger_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从Swagger数据中提取API接口"""
        interfaces = []
        paths = swagger_data.get('paths', {})
        servers = swagger_data.get('servers', [])
        base_url = servers[0].get('url', '') if servers else ''
        
        for path, path_item in paths.items():
            # 遍历所有HTTP方法
            for method, operation in path_item.items():
                if method.lower() in ['get', 'post', 'put', 'delete', 'patch', 'head', 'options']:
                    # 获取完整URL
                    full_url = f"{base_url}{path}" if base_url else path
                    
                    # 提取参数
                    parameters = operation.get('parameters', [])
                    params = {}
                    headers = {}
                    for param in parameters:
                        param_name = param.get('name', '')
                        param_in = param.get('in', '')
                        if param_in == 'query':
                            params[param_name] = param.get('schema', {}).get('default', '')
                        elif param_in == 'header':
                            headers[param_name] = param.get('schema', {}).get('default', '')
                    
                    # 提取请求体
                    request_body = operation.get('requestBody', {})
                    body = {}
                    if request_body:
                        content = request_body.get('content', {})
                        for content_type, schema_obj in content.items():
                            body_schema = schema_obj.get('schema', {})
                            body = {
                                'content_type': content_type,
                                'schema': body_schema
                            }
                            break
                    
                    # 提取响应schema
                    responses = operation.get('responses', {})
                    response_schema = {}
                    if '200' in responses or 'default' in responses:
                        success_response = responses.get('200') or responses.get('default', {})
                        response_content = success_response.get('content', {})
                        for content_type, schema_obj in response_content.items():
                            response_schema = {
                                'content_type': content_type,
                                'schema': schema_obj.get('schema', {})
                            }
                            break
                    
                    # 提取版本信息
                    version = self._extract_version_from_url(full_url)
                    
                    interface = {
                        'name': operation.get('summary') or operation.get('operationId') or f"{method.upper()} {path}",
                        'method': method.upper(),
                        'url': full_url,
                        'description': operation.get('description', ''),
                        'headers': headers,
                        'params': params,
                        'body': body,
                        'response_schema': response_schema,
                        'tags': operation.get('tags', []),
                        'deprecated': operation.get('deprecated', False),
                        'version': version
                    }
                    interfaces.append(interface)
        
        return interfaces
    
    def _extract_version_from_url(self, url: str) -> str:
        """从URL中提取版本信息"""
        if not url:
            return ""
        
        # 匹配常见的版本模式：/V0.1/, /V1.0/, /v1/, /api/v2/, /V0.1/等
        version_patterns = [
            r'/(V\d+\.\d+)/',      # /V0.1/, /V1.0/, /V2.3/
            r'/(v\d+\.\d+)/',      # /v0.1/, /v1.0/
            r'/v(\d+)/',           # /v1/, /v2/
            r'/api/v(\d+)/',       # /api/v1/, /api/v2/
            r'/version[_-]?(\d+\.\d+)/',  # /version-1.0/, /version_1.0/
        ]
        
        for pattern in version_patterns:
            match = re.search(pattern, url, re.IGNORECASE)
            if match:
                version = match.group(1)
                # 如果提取的是v开头，转换为V大写
                if version.startswith('v'):
                    version = 'V' + version[1:]
                elif not version.startswith('V'):
                    version = 'V' + version
                return version
        
        return ""
    
    def extract_api_interfaces(self, parsed_data: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从解析结果中提取API接口"""
        interfaces = []
        
        # 检查是否有直接的interfaces字段（Swagger、Postman、Apifox解析结果）
        if 'interfaces' in parsed_data and isinstance(parsed_data['interfaces'], list):
            interfaces = parsed_data['interfaces']
        
        # 检查raw_result中的interfaces
        elif 'raw_result' in parsed_data and isinstance(parsed_data['raw_result'], dict):
            raw = parsed_data['raw_result']
            if 'interfaces' in raw:
                interfaces = raw['interfaces']
            elif 'test_plans' in raw:
                for plan in raw['test_plans']:
                    for sample in plan.get('samples', []):
                        url = f"{sample.get('domain', '')}{sample.get('path', '')}"
                        interfaces.append({
                            'name': sample.get('name', ''),
                            'method': sample.get('method', 'GET'),
                            'url': url,
                            'description': 'JMeter测试采样器',
                            'version': self._extract_version_from_url(url)
                        })
        
        # 如果还没有接口，尝试从文本中提取（PDF、Markdown等）
        if not interfaces and 'text' in parsed_data:
            text_interfaces = self._extract_interfaces_from_text(parsed_data['text'])
            interfaces.extend(text_interfaces)
        
        # 为所有接口提取并添加版本信息
        for interface in interfaces:
            if 'version' not in interface or not interface.get('version'):
                url = interface.get('url', '') or interface.get('path', '')
                interface['version'] = self._extract_version_from_url(url)
        
        return interfaces
    
    def _extract_interfaces_from_text(self, text_list: List[str]) -> List[Dict[str, Any]]:
        """从文本中提取API接口信息（使用正则表达式）"""
        interfaces = []
        if not text_list:
            return interfaces
        
        # 合并所有文本
        full_text = '\n'.join(text_list)
        
        # 方法1: 匹配Markdown格式的接口描述（支持HTML标签）
        # 模式：数字. 接口名 ... URL: ... 服务: ...
        
        # 先清理HTML标签，但保留基本结构
        text_without_html = re.sub(r'<font[^>]*>', '', full_text)
        text_without_html = re.sub(r'</font>', '', text_without_html)
        
        # 使用finditer查找所有接口开始位置（数字. 开头）
        # 匹配模式：行首 + 可选#号 + 数字 + . + 接口名
        interface_start_pattern = re.compile(
            r'(?:^|\n)\s*(?:#+\s*)?(\d+)[\.、]\s*([^\n]+?)(?:\n|$)',
            re.MULTILINE
        )
        
        matches = list(interface_start_pattern.finditer(text_without_html))
        
        for i, match in enumerate(matches):
            num = match.group(1)
            name = match.group(2).strip()
            
            # 确定这个接口的结束位置（下一个接口开始或文档结束）
            start_pos = match.end()
            end_pos = matches[i+1].start() if i+1 < len(matches) else len(text_without_html)
            
            # 提取这个接口的部分内容
            section_content = text_without_html[start_pos:end_pos]
            
            # 提取URL（查找URL:后面的内容，支持各种格式）
            url_patterns = [
                r'URL[:：]\s*[`]?([^\s`\n\)]+)',  # URL: https://...
                r'url[:：]\s*[`]?([^\s`\n\)]+)',  # url: https://...
                r'https?://[^\s`\n\)]+',           # 直接匹配https://...
            ]
            
            url = None
            for pattern in url_patterns:
                url_match = re.search(pattern, section_content, re.IGNORECASE)
                if url_match:
                    url = url_match.group(1) if url_match.lastindex else url_match.group(0)
                    url = url.strip().rstrip('",\'')
                    if url.startswith('http'):
                        break
            
            # 提取服务名（查找服务:后面的内容）
            service_patterns = [
                r'服务[:：]\s*[`]?([^\s`\n\)]+)',
                r'service[:：]\s*[`]?([^\s`\n\)]+)',
            ]
            
            service = None
            for pattern in service_patterns:
                service_match = re.search(pattern, section_content, re.IGNORECASE)
                if service_match:
                    service = service_match.group(1).strip().rstrip('",\'')
                    break
            
            if url:  # 只要有URL就算有效接口
                # 推断HTTP方法（如果有请求体，通常是POST）
                method = 'POST'  # 默认POST（因为大部分接口都有请求体）
                # 检查是否有GET关键字在URL附近
                url_pos = section_content.find(url)
                context_start = max(0, url_pos - 100)
                context_end = min(len(section_content), url_pos + 100)
                context = section_content[context_start:context_end]
                if re.search(r'\bGET\b', context, re.IGNORECASE):
                    method = 'GET'
                
                # 提取版本信息
                version = self._extract_version_from_url(url)
                
                interfaces.append({
                    'name': name,
                    'method': method,
                    'url': url,
                    'service': service or '',
                    'description': f'{name} - {service}' if service else name,
                    'version': version
                })
        
        # 方法2: 如果方法1没找到，使用原有逻辑
        if not interfaces:
            # HTTP方法模式
            http_methods = ['GET', 'POST', 'PUT', 'DELETE', 'PATCH', 'HEAD', 'OPTIONS']
            
            # URL模式：http://... 或 https://... 或 /api/...
            url_pattern = re.compile(
                r'(?:'
                r'(?:http://|https://)[^\s\)`]+|'  # 完整URL（排除反引号）
                r'/api/[^\s\)`]+|'                  # 相对路径
                r'/v\d+/[^\s\)`]+|'                 # 版本化路径
                r'/[a-zA-Z][^\s\)`]+'               # 其他路径
                r')',
                re.IGNORECASE
            )
            
            # 查找接口模式
            lines = full_text.split('\n')
            
            for i, line in enumerate(lines):
                line_clean = line.strip()
                if not line_clean:
                    continue
                
                # 清理HTML标签
                line_clean = re.sub(r'<[^>]+>', '', line_clean)
                
                # 检查是否包含HTTP方法
                method = None
                for m in http_methods:
                    if re.search(rf'\b{m}\b', line_clean, re.IGNORECASE):
                        method = m.upper()
                        break
                
                # 查找URL（排除代码块中的内容，除非明确标记为URL）
                urls = url_pattern.findall(line_clean)
                
                # 如果找到方法和URL，创建接口
                if method and urls:
                    interface_name = line_clean.replace(method, '').strip()
                    # 移除URL和其他符号，保留接口名称
                    for url in urls:
                        interface_name = interface_name.replace(url, '').strip()
                    interface_name = re.sub(r'[^\w\s\u4e00-\u9fa5]', '', interface_name).strip()
                    
                    if not interface_name or len(interface_name) > 100:
                        # 尝试从上一行或下一行获取名称
                        for offset in [-1, 1]:
                            check_line_idx = i + offset
                            if 0 <= check_line_idx < len(lines):
                                check_line = re.sub(r'<[^>]+>', '', lines[check_line_idx]).strip()
                                if check_line and not url_pattern.search(check_line):
                                    interface_name = check_line[:50]
                                    break
                    
                    interfaces.append({
                        'name': interface_name or f'{method}接口',
                        'method': method,
                        'url': urls[0].rstrip('",\''),  # 清理末尾的引号
                        'description': line_clean[:200]
                    })
                
                # 如果只有URL没有方法，检查前后行是否有方法
                elif urls and not method:
                    # 检查前后5行
                    context_lines = []
                    for offset in range(max(0, i-5), min(len(lines), i+6)):
                        if offset != i:
                            clean_line = re.sub(r'<[^>]+>', '', lines[offset]).strip()
                            context_lines.append(clean_line)
                    
                    context_text = ' '.join(context_lines)
                    for m in http_methods:
                        if re.search(rf'\b{m}\b', context_text, re.IGNORECASE):
                            method = m.upper()
                            # 从上下文提取名称
                            interface_name = ''
                            for ctx_line in context_lines:
                                if ctx_line and not url_pattern.search(ctx_line):
                                    # 尝试提取标题格式的名称（数字. 名称）
                                    title_match = re.match(r'\d+[\.、]\s*(.+)', ctx_line)
                                    if title_match:
                                        interface_name = title_match.group(1).strip()[:50]
                                        break
                                    interface_name = ctx_line[:50]
                                    break
                            
                            interfaces.append({
                                'name': interface_name or f'{method}接口',
                                'method': method,
                                'url': urls[0].rstrip('",\''),  # 清理末尾的引号
                                'description': line_clean[:200]
                            })
                            break
        
        # 去重（基于URL、方法和服务名）
        seen = set()
        unique_interfaces = []
        for iface in interfaces:
            # 使用URL、方法和服务名作为唯一键
            service = iface.get('service', '')
            url = iface.get('url', '')
            method = iface.get('method', 'GET')
            
            # 清理URL中的引号
            url = url.rstrip('",\'').strip()
            
            key = (method, url, service) if service else (method, url)
            if key not in seen and url:
                seen.add(key)
                # 更新URL（清理后的）
                iface['url'] = url
                unique_interfaces.append(iface)
        
        return unique_interfaces
